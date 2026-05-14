"""Resume text extraction service for PDF and DOCX files."""

import io
import re
import fitz  # PyMuPDF
from docx import Document
from fastapi import UploadFile


async def extract_text(file: UploadFile) -> str:
    """Extract text content from an uploaded PDF or DOCX file.

    Args:
        file: The uploaded file (PDF or DOCX).

    Returns:
        Cleaned text content from the resume.

    Raises:
        ValueError: If the file type is not supported.
    """
    content = await file.read()
    filename = file.filename or ""

    if filename.lower().endswith(".pdf"):
        text = _extract_pdf(content)
    elif filename.lower().endswith(".docx"):
        text = _extract_docx(content)
    else:
        raise ValueError(
            f"Unsupported file type: {filename}. Only PDF and DOCX files are accepted."
        )

    return _clean_text(text)


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF."""
    doc = fitz.open(stream=content, filetype="pdf")
    pages = []
    for page in doc:
        pages.append(page.get_text("text"))
    doc.close()
    return "\n".join(pages)


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    doc = Document(io.BytesIO(content))
    paragraphs = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraphs.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def _clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    # Remove excessive whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    # Remove non-printable characters (keep newlines)
    text = re.sub(r"[^\S\n]+", " ", text)
    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)
    return text.strip()
