"""Service for compiling Resume Builder models into styled, ATS-compliant Microsoft Word (.docx) files."""

import io
import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


def generate_builder_docx(resume_data: dict) -> bytes:
    """Generate a clean, structured, and ATS-friendly DOCX from Resume Builder data."""
    doc = Document()

    # Margins: 0.75" top/bottom, 0.85" left/right for professional layout balance
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    # Color palette
    PRIMARY_COLOR = (37, 99, 235)  # Royal Blue (#2563eb)
    TEXT_COLOR = (17, 24, 39)      # Slate-900
    MUTED_COLOR = (107, 114, 128)  # Gray-500

    def add_section_header(title: str):
        # Add heading paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(title.upper())
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*PRIMARY_COLOR)

        # Add a clean bottom border line
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6") # 3/4 pt line weight
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "2563EB") # hex for Royal Blue
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_side_by_side_row(left_title: str, left_subtitle: str, right_text: str, right_subtext: str = ""):
        """Create a 2-column borderless table row to ensure perfectly aligned dates and titles."""
        tbl = doc.add_table(rows=1, cols=2)
        tbl.autofit = False
        
        # Set widths (6.8 inches total printable width)
        tbl.columns[0].width = Inches(5.1)
        tbl.columns[1].width = Inches(1.7)
        
        row = tbl.rows[0]
        
        # Left column (Title & Company/School)
        cell_l = row.cells[0]
        p_l = cell_l.paragraphs[0]
        p_l.paragraph_format.space_after = Pt(2)
        
        run_title = p_l.add_run(left_title)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(*TEXT_COLOR)
        
        if left_subtitle:
            run_sub = p_l.add_run(f"  |  {left_subtitle}")
            run_sub.font.bold = True
            run_sub.font.color.rgb = RGBColor(*PRIMARY_COLOR)
            
        # Right column (Dates & Location)
        cell_r = row.cells[1]
        p_r = cell_r.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_r.paragraph_format.space_after = Pt(2)
        
        run_r1 = p_r.add_run(right_text)
        run_r1.font.bold = True
        run_r1.font.color.rgb = RGBColor(*TEXT_COLOR)
        
        if right_subtext:
            p_r2 = cell_r.add_paragraph()
            p_r2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_r2.paragraph_format.space_after = Pt(2)
            run_r2 = p_r2.add_run(right_subtext)
            run_r2.font.size = Pt(9)
            run_r2.font.italic = True
            run_r2.font.color.rgb = RGBColor(*MUTED_COLOR)

    # ── Header: Personal Info ──
    personal_info = resume_data.get("personalInfo", {})
    name = personal_info.get("fullName") or "Your Name"
    
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(4)
    run_name = name_para.add_run(name)
    run_name.font.size = Pt(20)
    run_name.font.bold = True
    run_name.font.color.rgb = RGBColor(*TEXT_COLOR)

    if personal_info.get("professionalTitle"):
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(4)
        run_title = title_para.add_run(personal_info.get("professionalTitle").upper())
        run_title.font.size = Pt(10.5)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(*PRIMARY_COLOR)

    # Contacts list
    contact_parts = [
        personal_info.get("email"),
        personal_info.get("phone"),
        personal_info.get("location"),
        personal_info.get("linkedin") and f"linkedin.com/in/{personal_info.get('linkedin').split('/').pop()}",
        personal_info.get("github") and f"github.com/{personal_info.get('github').split('/').pop()}",
        personal_info.get("portfolioWebsite")
    ]
    contact_parts = [c for c in contact_parts if c]
    
    if contact_parts:
        contact_para = doc.add_paragraph("   ·   ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(12)
        for run in contact_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(*MUTED_COLOR)

    # ── Professional Summary ──
    summary = personal_info.get("professionalSummary", "")
    if summary:
        add_section_header("Professional Summary")
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15

    # ── Work Experience ──
    experience = resume_data.get("experience", [])
    if experience:
        add_section_header("Experience")
        for exp in experience:
            dates = f"{exp.get('startDate')} – {exp.get('endDate') or ('Present' if exp.get('current') else '')}"
            add_side_by_side_row(
                left_title=exp.get("jobTitle", ""),
                left_subtitle=exp.get("company", ""),
                right_text=dates,
                right_subtext=exp.get("location", "")
            )
            
            # Responsibilities list
            resp_str = exp.get("responsibilities", "")
            if resp_str:
                for line in resp_str.split("\n"):
                    if line.strip():
                        bp = doc.add_paragraph(style="List Bullet")
                        bp.paragraph_format.left_indent = Inches(0.2)
                        bp.paragraph_format.space_after = Pt(2)
                        bp.add_run(line.strip())
            
            # Key achievements
            ach = exp.get("achievements", "")
            if ach:
                ach_p = doc.add_paragraph()
                ach_p.paragraph_format.left_indent = Inches(0.2)
                ach_p.paragraph_format.space_after = Pt(4)
                bold_run = ach_p.add_run("Key Achievement: ")
                bold_run.font.bold = True
                bold_run.font.size = Pt(9)
                bold_run.font.color.rgb = RGBColor(*PRIMARY_COLOR)
                
                italic_run = ach_p.add_run(ach)
                italic_run.font.italic = True
                italic_run.font.size = Pt(9)
                italic_run.font.color.rgb = RGBColor(*TEXT_COLOR)

            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Projects ──
    projects = resume_data.get("projects", [])
    if projects:
        add_section_header("Projects")
        for proj in projects:
            links = []
            if proj.get("githubLink"):
                links.append("GitHub")
            if proj.get("liveDemoLink"):
                links.append("Live Demo")
            
            right_label = "  |  ".join(links) if links else ""
            add_side_by_side_row(
                left_title=proj.get("projectName", ""),
                left_subtitle="",
                right_text=right_label
            )

            # Technologies tag list
            techs = proj.get("technologies", [])
            if techs:
                tech_p = doc.add_paragraph()
                tech_p.paragraph_format.left_indent = Inches(0.15)
                tech_p.paragraph_format.space_after = Pt(2)
                bold_run = tech_p.add_run("Technologies: ")
                bold_run.font.bold = True
                bold_run.font.size = Pt(8.5)
                bold_run.font.color.rgb = RGBColor(*MUTED_COLOR)
                
                tech_run = tech_p.add_run(", ".join(techs))
                tech_run.font.size = Pt(8.5)
                tech_run.font.color.rgb = RGBColor(*TEXT_COLOR)

            # Description
            desc = proj.get("description", "")
            if desc:
                dp = doc.add_paragraph(desc)
                dp.paragraph_format.left_indent = Inches(0.15)
                dp.paragraph_format.space_after = Pt(4)

            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Education ──
    education = resume_data.get("education", [])
    if education:
        add_section_header("Education")
        for edu in education:
            dates = f"{edu.get('startDate')} – {edu.get('endDate')}"
            gpa_sub = f"GPA: {edu.get('gpa')}" if edu.get("gpa") else ""
            add_side_by_side_row(
                left_title=edu.get("degree", ""),
                left_subtitle=edu.get("institution", ""),
                right_text=dates,
                right_subtext=gpa_sub
            )
            
            desc = edu.get("description", "")
            if desc:
                edup = doc.add_paragraph(desc)
                edup.paragraph_format.left_indent = Inches(0.15)
                edup.paragraph_format.space_after = Pt(4)

            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Skills & Credentials ──
    skills = resume_data.get("skills", {})
    has_skills = any(skills.values())
    if has_skills:
        add_section_header("Skills & Credentials")
        
        cats = [
            ("Technical Skills", skills.get("technical", [])),
            ("Soft Skills", skills.get("soft", [])),
            ("Languages", skills.get("languages", [])),
            ("Certifications", skills.get("certifications", []))
        ]
        
        for label, items in cats:
            if items:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                bold_run = p.add_run(f"{label}: ")
                bold_run.font.bold = True
                bold_run.font.color.rgb = RGBColor(*TEXT_COLOR)
                
                val_run = p.add_run(", ".join(items))
                val_run.font.color.rgb = RGBColor(*TEXT_COLOR)

    # Serialize
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
