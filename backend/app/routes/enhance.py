"""Enhancement API routes — AI resume rewriting and DOCX download."""

import io
import logging
from typing import Optional

from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse

from app.config import settings
from app.models.schemas import (
    EnhanceRequest, EnhanceResponse, EnhancedResumeRecord, ErrorResponse
)
from app.services.enhancer import enhance_resume, VALID_MODES
from app.services.auth import get_current_user
from app.database import get_db
from app.routes.upload import get_guest_store

logger = logging.getLogger(__name__)

router = APIRouter()

# In-memory store for guest-enhanced resumes
_guest_enhance_store: dict[str, EnhancedResumeRecord] = {}


def _lookup_analysis_text(analysis_id: str, user: Optional[dict]) -> str:
    """Retrieve resume_text from Firestore (auth) or guest store."""
    if user:
        db = get_db()
        doc = db.collection("analyses").document(analysis_id).get()
        if not doc.exists:
            raise HTTPException(status_code=404, detail="Analysis not found.")
        data = doc.to_dict()
        resume_text = data.get("resume_text", "")
    else:
        guest_store = get_guest_store()
        record = guest_store.get(analysis_id)
        if not record:
            raise HTTPException(
                status_code=404,
                detail="Analysis not found. Guest analyses are session-only.",
            )
        resume_text = record.resume_text

    if not resume_text.strip():
        raise HTTPException(
            status_code=400,
            detail="No resume text found for this analysis. Cannot enhance.",
        )
    return resume_text


# ─── POST /api/enhance ──────────────────────────────────────────────────────────

@router.post(
    "/enhance",
    response_model=EnhanceResponse,
    responses={400: {"model": ErrorResponse}, 404: {"model": ErrorResponse}, 500: {"model": ErrorResponse}},
    summary="Enhance a resume with AI",
    description=(
        "Rewrites and improves a previously analyzed resume using AI. "
        "Choose from four modes: professional, technical, executive, fresher. "
        "Returns before/after content for every section."
    ),
)
async def enhance_resume_endpoint(
    request: EnhanceRequest,
    user: Optional[dict] = Depends(get_current_user),
):
    """AI resume enhancement endpoint.

    Requires a valid analysis_id from a previous /upload call.
    Uses the same AI provider (Groq or Ollama) as the analysis pipeline.
    """
    # Validate mode
    if request.mode not in VALID_MODES:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid mode '{request.mode}'. Valid modes: {', '.join(VALID_MODES)}",
        )

    # Fetch original resume text
    resume_text = _lookup_analysis_text(request.analysis_id, user)

    # Run AI enhancement
    try:
        logger.info(
            f"Starting enhancement — id={request.analysis_id}, "
            f"mode={request.mode}, user={'auth' if user else 'guest'}"
        )
        enhanced = await enhance_resume(
            resume_text=resume_text,
            mode=request.mode,
            job_description=request.job_description,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        logger.exception(f"Unexpected enhancement error for {request.analysis_id}")
        raise HTTPException(status_code=500, detail=f"Enhancement failed: {e}")

    # Persist record
    record = EnhancedResumeRecord(
        analysis_id=request.analysis_id,
        mode=request.mode,
        enhanced_resume=enhanced,
        job_description=request.job_description,
        user_id=user["uid"] if user else None,
    )

    if user:
        db = get_db()
        doc_data = record.model_dump()
        doc_data["enhanced_at"] = record.enhanced_at.isoformat()
        db.collection("enhancements").document(record.id).set(doc_data)
        logger.info(f"Enhancement saved to Firestore: {record.id}")
    else:
        _guest_enhance_store[record.id] = record
        logger.info(f"Guest enhancement stored in-memory: {record.id}")

    return EnhanceResponse(
        id=record.id,
        analysis_id=record.analysis_id,
        mode=record.mode,
        enhanced_resume=enhanced,
        enhanced_at=record.enhanced_at,
    )


# ─── GET /api/enhance/{id}/download/docx ────────────────────────────────────────

@router.get(
    "/enhance/{enhancement_id}/download/docx",
    responses={404: {"model": ErrorResponse}, 500: {"model": ErrorResponse}},
    summary="Download enhanced resume as DOCX",
)
async def download_enhanced_docx(
    enhancement_id: str,
    social_links_json: Optional[str] = None,
    display_mode: str = "compact",
    user: Optional[dict] = Depends(get_current_user),
):
    """Stream a DOCX file of the AI-enhanced resume."""
    # Look up the enhancement record
    if user:
        db = get_db()
        doc = db.collection("enhancements").document(enhancement_id).get()
        if not doc.exists:
            raise HTTPException(status_code=404, detail="Enhancement not found.")
        data = doc.to_dict()
        record = EnhancedResumeRecord.model_validate(data)
    else:
        record = _guest_enhance_store.get(enhancement_id)
        if not record:
            raise HTTPException(status_code=404, detail="Enhancement not found.")

    # Fetch social links for authenticated user or fallback to param
    social_links = []
    if user:
        try:
            db = get_db()
            social_doc = db.collection("social_links").document(user["uid"]).get()
            if social_doc.exists:
                social_data = social_doc.to_dict()
                social_links = social_data.get("links", [])
                display_mode = social_data.get("display_mode", display_mode)
        except Exception as e:
            logger.warning(f"Failed to fetch social links from Firestore for DOCX: {e}")
            
    if not social_links and social_links_json:
        try:
            import json
            social_links = json.loads(social_links_json)
        except Exception as e:
            logger.warning(f"Failed to parse social_links_json: {e}")

    try:
        docx_bytes = _generate_docx(record.enhanced_resume, social_links, display_mode)
    except Exception as e:
        logger.exception("DOCX generation failed")
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {e}")

    filename = f"enhanced_resume_{record.mode}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─── DOCX Generation ─────────────────────────────────────────────────────────────

def _generate_docx(enhanced: "EnhancedResume", social_links: list = None, display_mode: str = "compact") -> bytes:
    """Generate a styled DOCX from an EnhancedResume object."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # --- Page margins ---
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # --- Styles ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    def add_heading(text: str, level: int = 1, color: tuple = (79, 70, 229)):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
            run.font.bold = True
        return p

    def add_section_divider():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "4F46E5")
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(2)

    # ── Header: Name + Contact ──
    name = enhanced.candidate_name or "Candidate"
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(name)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 31, 71)

    if enhanced.contact_info:
        contact_para = doc.add_paragraph(enhanced.contact_info)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(107, 114, 128)

    # ── Social Links ──
    active_social = [l for l in (social_links or []) if l.get("is_enabled", True)]
    if active_social:
        import docx.opc.constants
        
        social_para = doc.add_paragraph()
        social_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        social_para.paragraph_format.space_before = Pt(2)
        social_para.paragraph_format.space_after = Pt(2)
        
        def add_docx_hyperlink(paragraph, url, text):
            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            c = OxmlElement('w:color')
            c.set(qn('w:val'), '4F46E5')
            rPr.append(c)
            
            u = OxmlElement('w:underline')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '19')
            rPr.append(sz)
            
            new_run.append(rPr)
            text_node = OxmlElement('w:t')
            text_node.text = text
            new_run.append(text_node)
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
            return hyperlink

        for idx, item in enumerate(active_social):
            platform_name = item.get("platform", "custom").capitalize()
            clean_url = item.get("url", "").replace("https://", "").replace("http://", "").replace("www.", "")
            
            if display_mode == "ats_safe":
                label_text = clean_url
            elif display_mode == "icon_only":
                label_text = f"[{platform_name}]"
            else:
                label = item.get("label") or platform_name
                label_text = f"{label} ({clean_url})" if display_mode == "expanded" else label
                
            add_docx_hyperlink(social_para, item.get("url", ""), label_text)
            
            if idx < len(active_social) - 1:
                sep_run = social_para.add_run("  ·  ")
                sep_run.font.size = Pt(9.5)
                sep_run.font.color.rgb = RGBColor(156, 163, 175)

    doc.add_paragraph()  # spacer





    # ── Professional Summary ──
    summary = enhanced.professional_summary
    if summary.enhanced:
        add_heading("PROFESSIONAL SUMMARY", level=2, color=(79, 70, 229))
        add_section_divider()
        p = doc.add_paragraph(summary.enhanced)
        p.paragraph_format.space_after = Pt(8)

    # ── Work Experience ──
    if enhanced.experience_sections:
        add_heading("WORK EXPERIENCE", level=2, color=(79, 70, 229))
        add_section_divider()
        for exp in enhanced.experience_sections:
            # Company + Title row
            exp_para = doc.add_paragraph()
            title_run = exp_para.add_run(f"{exp.title}" if exp.title else "")
            title_run.font.bold = True
            title_run.font.size = Pt(11)
            if exp.company:
                sep_run = exp_para.add_run(f"  |  {exp.company}")
                sep_run.font.color.rgb = RGBColor(79, 70, 229)
                sep_run.font.size = Pt(10.5)
            if exp.duration:
                dur_run = exp_para.add_run(f"  ·  {exp.duration}")
                dur_run.font.size = Pt(9.5)
                dur_run.font.color.rgb = RGBColor(107, 114, 128)

            # Bullets
            bullets = exp.enhanced_bullets or exp.original_bullets
            for bullet in bullets:
                bp = doc.add_paragraph(style="List Bullet")
                bp.text = bullet
                bp.paragraph_format.left_indent = Inches(0.2)
                bp.paragraph_format.space_after = Pt(2)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Projects ──
    if enhanced.projects_sections:
        add_heading("PROJECTS", level=2, color=(79, 70, 229))
        add_section_divider()
        for proj in enhanced.projects_sections:
            proj_para = doc.add_paragraph()
            name_run = proj_para.add_run(proj.name or "Project")
            name_run.font.bold = True
            name_run.font.size = Pt(11)
            desc = proj.enhanced or proj.original
            if desc:
                dp = doc.add_paragraph(desc)
                dp.paragraph_format.left_indent = Inches(0.2)
                dp.paragraph_format.space_after = Pt(4)

    # ── Education ──
    if enhanced.education_sections:
        add_heading("EDUCATION", level=2, color=(79, 70, 229))
        add_section_divider()
        for edu in enhanced.education_sections:
            edu_para = doc.add_paragraph()
            deg_run = edu_para.add_run(edu.degree or "")
            deg_run.font.bold = True
            if edu.institution:
                inst_run = edu_para.add_run(f"  —  {edu.institution}")
                inst_run.font.color.rgb = RGBColor(79, 70, 229)
            if edu.year:
                yr_run = edu_para.add_run(f"  ({edu.year})")
                yr_run.font.color.rgb = RGBColor(107, 114, 128)
            enhanced_text = edu.enhanced or edu.original
            if enhanced_text and enhanced_text != edu.degree:
                extra = doc.add_paragraph(enhanced_text)
                extra.paragraph_format.left_indent = Inches(0.2)
                extra.paragraph_format.space_after = Pt(4)

    # ── Skills ──
    skills = enhanced.skills_section.enhanced or enhanced.skills_section.original
    if skills:
        add_heading("SKILLS", level=2, color=(79, 70, 229))
        add_section_divider()
        # Group into rows of 5
        chunks = [skills[i:i+5] for i in range(0, len(skills), 5)]
        for chunk in chunks:
            p = doc.add_paragraph("  ·  ".join(chunk))
            p.paragraph_format.space_after = Pt(2)

    # ── Achievements ──
    achievements = enhanced.achievements_section.enhanced or enhanced.achievements_section.original
    if achievements:
        add_heading("ACHIEVEMENTS", level=2, color=(79, 70, 229))
        add_section_divider()
        for ach in achievements:
            ap = doc.add_paragraph(style="List Bullet")
            ap.text = ach
            ap.paragraph_format.space_after = Pt(2)

    # ── Footer ──
    doc.add_paragraph()
    footer_para = doc.add_paragraph("✨ Enhanced by Smart Resume Analyzer AI")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(156, 163, 175)

    # Serialize to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
